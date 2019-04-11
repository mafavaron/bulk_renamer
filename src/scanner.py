import sys
import os
from docx import Document

new_preamble = {
    'CA': 'E1844-DT-SF-0072, “Analisi dei siti per il posizionamento delle stazioni di misura” – Volume 2 – Stralcio CAGLIARI – ',
    'CI': 'E1844-DT-SF-0073, “Analisi dei siti per il posizionamento delle stazioni di misura” – Volume 2 – Stralcio CARBONIA-IGLESIAS – ',
    'VS': 'E1844-DT-SF-0074, “Analisi dei siti per il posizionamento delle stazioni di misura” – Volume 2 – Stralcio MEDIO CAMPIDANO – ',
    'NU': 'E1844-DT-SF-0075, “Analisi dei siti per il posizionamento delle stazioni di misura” – Volume 2 – Stralcio NUORO – ',
    'OG': 'E1844-DT-SF-0076, “Analisi dei siti per il posizionamento delle stazioni di misura” – Volume 2 – Stralcio OGLIASTRA – ',
    'OT': 'E1844-DT-SF-0077, “Analisi dei siti per il posizionamento delle stazioni di misura” – Volume 2 – Stralcio OLBIA-TEMPIO PAUSANIA – ',
    'OR': 'E1844-DT-SF-0078, “Analisi dei siti per il posizionamento delle stazioni di misura” – Volume 2 – Stralcio ORISTANO – ',
    'SS': 'E1844-DT-SF-0079, “Analisi dei siti per il posizionamento delle stazioni di misura” – Volume 2 – Stralcio SASSARI – '
}


def textReplace(current_doc_name, new_doc_name, site, ru, prov, report_file):

    # Form input and output directory names
    type = ru[-3:-2]
    in_path  = os.path.join(".", "Schede_Stazione_Vecchie", "Tipo_%s" % type, prov, current_doc_name)
    out_dir = os.path.join(".", "Schede_Stazione_Nuove", "Tipo_%s" % type, prov)
    if not os.path.exists(out_dir):
        os.makedirs(out_dir)
    out_path = os.path.join(out_dir, new_doc_name)

    # Access document, to perform the renames due
    print("Old file: %s" % current_doc_name)
    success = True
    try:
        document = Document(in_path)
    except Exception as e:
        success = False

    # Iterate over document header contents
    for sect in range(len(document.sections)):
        section = document.sections[sect]

        # Header: Change file name to be consistent with actual
        for par in section.header.paragraphs:
            preamble   = new_preamble[prov]
            prefix     = "Allegato \""
            postfix    = "\""
            new_name   = os.path.basename(new_doc_name).split(".")[0]
            new_para   = preamble + prefix + new_name + postfix
            par.text   = new_para

    # Find station name
    for par in document.paragraphs:
        if "Stazione" in par.text:
            blocks = par.text.split("“")
            if len(blocks) >= 2:
                parts = blocks[1].split("”")
                if len(parts) >= 2:
                    par.text = blocks[0] + "“" + site + "”" + parts[1]

    # Table on initial page:
    if len(document.tables) <= 0:
        success = False
    else:
        row_num = -1
        first_table = document.tables[0]
        for row in first_table.rows:
            row_num += 1
            if len(row.cells) >= 2:
                title = row.cells[0].text
                if "sopralluogo" in title:
                    row.cells[1].text = new_doc_name.replace("docx", "pdf")

    # Save new document
    if success:
        document.save(out_path)

    # Save data
    if success:
        report_file.write("Old file %s successfully processed.\n" % current_doc_name)
    else:
        report_file.write("Old file %s <<<<<<<<<< Error >>>>>>>>>>>>>>>>\n" % current_doc_name)


if __name__ == "__main__":

    if len(sys.argv) != 2:
        print("scanner.py - Bulk renamer for DOCX file contents")
        print()
        print("Usage:")
        print()
        print("  python scanner.py <Input_File>")
        print()
        print("Written by: Mauri Favaron")
        sys.exit(1)
    in_file     = sys.argv[1]

    # Scan input file
    prov     = []
    prov_id  = []
    ru       = []
    new_alt  = []
    old_alt  = []
    site     = []
    old_file = []
    new_file = []
    f = open(in_file, "r")
    in_lines = f.readlines()
    f.close()
    for line_idx in range(1, len(in_lines)):
        blocks = in_lines[line_idx].split(";")
        prov.append(blocks[0])
        prov_id.append(int(blocks[1]))
        ru.append(blocks[2])
        new_alt.append(blocks[3])
        old_alt.append(blocks[4])
        site.append(blocks[5])
        old_file.append(blocks[6])
        new_file.append(blocks[7])

    # Process data
    f = open("./processing.dat", "w")
    for idx in range(len(old_file)):
        textReplace(old_file[idx], new_file[idx], site[idx], ru[idx], prov[idx], f)
    f.close()
